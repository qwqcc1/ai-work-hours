"""Generate weekly plan report as Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── global style ──
style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * (level + 1))


def bullet2(text):
    bullet(text, level=1)


def set_cell(cell, text, bold=False, align_center=True, font_size=9, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align_center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tblPr.append(borders)


def shade_row(row, color="4472C4"):
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_border(table)
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, font_size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade_row(table.rows[0], "4472C4")
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            set_cell(table.rows[r].cells[c], val, bold=(c == 0), font_size=9, align_center=(c != 1))
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI工时管理系统\n周报计划")
run.font.size = Pt(28)
run.bold = True
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("2026年6月9日  ·  已完成功能 & 待优化 & 两周计划")
run.font.size = Pt(12)
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════
# PART 1: COMPLETED FEATURES (EXPANDED)
# ═══════════════════════════════════════════
heading("第一部分：已完成功能", level=1)

# ── 1.1 用户认证系统 ──
heading("1.1 用户认证系统（8 个 API）", level=2)

para("1. 用户注册（POST /api/register）", bold=True)
bullet("用户填写用户名、密码（至少6位）、真实姓名、手机号即可自助注册")
bullet("Pydantic 数据校验：用户名 1-50 字符，密码至少 6 位")
bullet("后端查重：用户名已存在则返回 400「用户名已存在」")
bullet("注册即登录：成功后自动签发 JWT Token，创建会话记录，跳转仪表盘")
bullet("前端表单有原生 HTML5 校验（minlength、required），双重保障")

para("2. 用户登录（POST /api/login）", bold=True)
bullet("用户名 + 密码验证，支持 device_type（pc/mobile）和 device_name 记录")
bullet("密码使用 bcrypt 哈希比对，不可逆，安全性高")
bullet("登录成功返回 JWT Token + 用户信息（id, username, real_name, role）")
bullet("自动创建 UserSession 记录，跟踪登录设备、IP、活跃时间")
bullet("区分错误类型：401「用户名或密码错误」、403「账号已被禁用」")
bullet("Token 有效期 24 小时，支持多端同时在线互不影响")
bullet("前端登录按钮防重复提交（提交中禁用按钮）")

para("3. 多端会话管理（GET /api/sessions 等 4 个端点）", bold=True)
bullet("查看所有活跃会话：设备类型（PC/手机）、设备名、IP 地址、最后活跃时间")
bullet("单独登出指定设备（POST /api/logout-device）：需本人或管理员权限")
bullet("一键登出所有设备（POST /api/logout-all）")
bullet("登出当前会话（POST /api/logout）：清除 Token，标记会话失效")
bullet("会话最后活跃时间在每次 API 请求时自动更新（get_current_user 中间件）")

para("4. 个人设置（PUT /api/profile）", bold=True)
bullet("可修改真实姓名、手机号、密码（密码可选，留空则不修改）")
bullet("前端 Modal 弹窗编辑，保存后实时更新导航栏显示名和 localStorage 缓存")

para("5. 角色权限体系", bold=True)
bullet("三级角色：admin（管理员）、manager（经理）、member（普通成员）")
bullet("admin：可访问后台管理、查看所有数据、管理用户/项目/会话")
bullet("manager：可创建项目，其他同 member")
bullet("member：仅看自己的工时和任务，不可管理用户和项目")
bullet("所有敏感接口均有 role 校验，非 admin 返回 403")

doc.add_paragraph()

# ── 1.2 工时管理 ──
heading("1.2 工时管理（4 个 API）", level=2)

para("1. 工时录入（POST /api/work-entries）", bold=True)
bullet("8 个字段：项目、工时（0.5-24h，步长 0.5）、工时类型、日期、地区、项目番号、完成情况、工作描述")
bullet("Pydantic 严格校验：hours > 0 且 ≤ 24，所有必填字段检查")
bullet("自动捕获 User-Agent 作为设备来源（PC/手机浏览器标识），截断至 200 字符")
bullet("工时类型下拉 8 个选项：日常运维、故障处理、巡检监控、变更操作、专项工作、会议培训、技术支持、其他")
bullet("提交成功后自动清除除日期和项目外的表单字段，方便连续录入")
bullet("防重复提交：按钮显示「提交中...」并禁用")

para("2. 工时列表（GET /api/work-entries）", bold=True)
bullet("支持多条件筛选：project_id、start_date、end_date、limit（默认 100，最大 10000）")
bullet("权限隔离：非 admin 只能看到自己的记录，admin 可查看所有")
bullet("按 entry_date 降序排列，最新记录在前")
bullet("返回关联数据：user_name、project_name、project_category 一次查询带出")
bullet("前端展示：项目名带颜色标签（方舱=绿色，运维项目=蓝色），描述超过 20 字截断")

para("3. 工时编辑与删除", bold=True)
bullet("编辑（PUT /api/work-entries/{id}）：本人或管理员可修改任意字段")
bullet("删除（DELETE /api/work-entries/{id}）：软删除（is_active=False），数据可恢复")
bullet("权限校验：非本人且非 admin 返回 403「无权修改/删除他人的记录」")
bullet("前端删除有 Modal 确认弹窗「此操作不可撤销」，确认后才执行")

doc.add_paragraph()

# ── 1.3 仪表盘与统计 ──
heading("1.3 仪表盘与汇总统计", level=2)

para("1. 仪表盘（GET /dashboard）", bold=True)
bullet("4 个统计卡片：今日工时（蓝色）、本周工时+日均（绿色）、本月工时+日均（橙色）、本月记录数+活跃项目数（紫色）")
bullet("每张卡片带 hover 动效（顶部色条渐显），数据空时显示「--」")
bullet("周趋势折线图（自定义 Canvas 图表库）：展示周一至周日每天工时，带渐变面积填充，数据点白底圆点标注")
bullet("项目分布环形图：当月各项目工时占比，中心显示总工时，下方图例带色点")
bullet("最近 10 条工时记录表：项目名带颜色标签，描述超过 30 字截断，空态引导「录入第一条工时」")
bullet("快捷导出按钮：一键导出当月 Excel 报表")
bullet("所有图表主题自适应：深色模式下自动切换配色，通过 getComputedStyle 读取 CSS 变量")

para("2. 汇总统计（POST /api/summary）", bold=True)
bullet("支持三维度：按日 / 按周（ISO 周，如 2026W23）/ 按月 汇总")
bullet("返回每组的 total_hours、entries_count、by_project 细分、by_user 细分")
bullet("筛选条件：日期范围（默认近 30 天）、项目、用户（仅 admin）、项目分类")
bullet("前端展示：汇总结果表格 + 顶部总计栏（大字显示累计工时）")
bullet("支持按项目分类（方舱/运维项目）过滤，在 Python 内存中筛选")
bullet("数据为空时显示「无匹配数据」空态，初始态显示「请选择条件后点击查询」")

para("3. 自定义 Canvas 图表库（SimpleChart）", bold=True)
bullet("纯手写 Canvas 图表库，支持柱状图、折线图、环形图三种类型")
bullet("柱状图：渐变填充、顶部圆角（roundRect）、数值标签、自适应 Y 轴")
bullet("折线图：渐变面积填充、数据点白底圆点、平滑连接线（rounded joins）")
bullet("环形图：从顶部 -PI/2 起始、60% 内径、中心文字、8 色调色板")
bullet("支持 Retina 高清屏（devicePixelRatio 适配）")
bullet("窗口 resize 时自动重绘，主题切换时实时响应")

doc.add_paragraph()

# ── 1.4 导出功能 ──
heading("1.4 Excel 导出系统", level=2)

para("1. 自由格式导出（POST /api/export）", bold=True)
bullet("生成双工作表 Excel：工时数据（13 列详细记录）+ 汇总统计（总工时、记录数、导出日期）")
bullet("蓝色表头（#4472C4），白色加粗微软雅黑字体，所有单元格居中+细边框")
bullet("13 列：序号、人员、项目名称、分类、项目番号、完成情况、地区、工时类型、工时(h)、日期、设备来源、工作描述、录入时间")
bullet("固定列宽优化排版，限 10000 条记录")
bullet("以字节流返回，Content-Disposition 附件下载，文件名含日期")

para("2. 模板月度导出（POST /api/export-monthly-template）", bold=True)
bullet("基于年度 Excel 模板（template.xlsx）填充指定月份工时到对应日期列")
bullet("自动扫描模板表头：正则匹配月份标题（如「6 月份工时统计」），解析第 3 行日号建立列映射")
bullet("按（用户, 项目番号）分组，一人多项目生成多行")
bullet("自动扩行/删行：用户数超过模板行数时自动插入行，不够则删除空行")
bullet("智能公式保留：行插入/删除后用正则修复所有 SUM/引用公式中的行号")
bullet("保有工时行（最后一行汇总）始终保留且永远在最底部")
bullet("合并单元格处理：先解合并再操作行，避免 openpyxl 报错")
bullet("完成情况取当月最多出现的那条（最高频状态）")
bullet("导出前解合并、操作后公式修复、BytesIO 内存流返回，不修改原模板文件")

doc.add_paragraph()

# ── 1.5 任务管理 ──
heading("1.5 任务管理系统（6 个 API）", level=2)

para("1. 发布任务（POST /api/tasks）", bold=True)
bullet("填写标题（必填，最多 200 字）、所属项目、开始/结束时间、地点、描述")
bullet("弹窗多选执行人：通过 /api/users/simple 获取用户列表，弹窗展示复选框")
bullet("选中的执行人以蓝色标签（chip）展示在表单上，可点击 × 移除")
bullet("前端校验：至少选一位执行人，否则 toast 提示")
bullet("开始时间默认当前，结束时间默认 +2 小时")

para("2. 任务列表", bold=True)
bullet("「我发布的」视图：按状态筛选（pending/in_progress/completed/cancelled），展示执行人名单和状态标签")
bullet("「分配给我的」视图：展示发布人、项目、时间、地点，按任务分配状态筛选")
bullet("状态颜色区分：待处理=橙色、进行中/已接受=蓝色、已完成=绿色、已拒绝/已取消=灰色")
bullet("三 Tab 切换界面（我发布的 / 分配给我的 / 发布新任务），URL hash 联动")

para("3. 任务操作（PUT /api/tasks/{id}/assign）", bold=True)
bullet("三种操作：accept（接受）→ 记录 accepted_at、reject（拒绝）、done（完成）→ 记录 completed_at")
bullet("自动任务状态联动：全部执行人 done → 任务 completed，有人 accepted/done → 任务 in_progress")
bullet("发布者或管理员可取消任务（DELETE → 状态改为 cancelled）")
bullet("前端动态按钮：待接收显示「接受」「拒绝」，已接受显示「完成」，已完成/已拒绝显示状态文字")

doc.add_paragraph()

# ── 1.6 后台管理 ──
heading("1.6 后台管理系统", level=2)

para("1. 用户管理（/admin#users）", bold=True)
bullet("管理员可创建用户（用户名、密码、姓名、电话、角色），用户名查重")
bullet("编辑用户：Modal 弹窗修改姓名、电话、角色（admin/manager/member）、状态（启用/禁用）")
bullet("删除用户：软删除（is_active=False），禁用后无法登录，数据保留")
bullet("空密码不修改密码（编辑时密码留空则跳过）")

para("2. 项目管理（/admin#projects）", bold=True)
bullet("创建项目：名称、分类（运维项目/方舱）、描述")
bullet("编辑项目：Modal 弹窗改任意字段")
bullet("删除项目：软删除，已有工时记录不受影响")
bullet("项目下拉列表仅显示 is_active=True 的项目")

para("3. 人员查看（/users）", bold=True)
bullet("全部人员列表：用户名、姓名、角色、手机号、状态、注册时间")
bullet("实时搜索过滤：支持按用户名/手机号/姓名模糊匹配（大小写不敏感）")
bullet("角色彩色标签（admin=红、manager=橙、member=灰）、状态标识（启用=绿、禁用=红）")
bullet("非 admin 访问自动跳转仪表盘")

para("4. 在线会话管理", bold=True)
bullet("查看所有活跃会话：设备类型图标（手机/PC）、设备名、IP、最后活跃时间")
bullet("单设备强制登出，一键登出所有设备前有二次确认")

doc.add_paragraph()

# ── 1.7 数据模型 ──
heading("1.7 数据库设计", level=2)

para("6 张数据表，SQLAlchemy ORM 管理，启动自动建表+迁移：", bold=True)
bullet("users：id, username(索引唯一), password_hash(bcrypt), real_name, role(admin/manager/member), phone, is_active, created_at")
bullet("projects：id, name, category(方舱/运维项目), description, is_active, created_at")
bullet("work_entries：id, user_id(FK), project_id(FK), hours, work_type(8种), description, entry_date(索引), device_info, location, project_number, completion_status, is_active, created_at, updated_at")
bullet("user_sessions：id, user_id(FK), token(索引), device_type, device_name, ip_address, is_active, last_activity")
bullet("tasks：id, publisher_id(FK), project_id(FK), title, description, start_time, end_time, location, status(pending/in_progress/completed/cancelled)")
bullet("task_assignments：id, task_id(FK), user_id(FK), status(pending/accepted/rejected/done), assigned_at, accepted_at, completed_at, remark")
bullet("所有删除操作均为软删除（is_active=False / status=cancelled），数据可追溯恢复")
bullet("启动时自动迁移：检测 work_entries 表缺少 project_number/completion_status 列则 ALTER TABLE 补全")

doc.add_paragraph()

# ── 1.8 前端体验 ──
heading("1.8 前端交互与 UI 系统", level=2)

para("1. 导航系统", bold=True)
bullet("固定顶部导航栏，毛玻璃效果（backdrop-filter: blur(20px)），三段式布局：品牌+链接+用户")
bullet("下拉菜单：「工作总结」（工时表、月总结、绩效评价、目标设定）、「任务管理」（查看、发布）、「系统管理」（仅管理员可见）")
bullet("下拉面板使用 Web Animations API 实现淡入淡出（250ms 入场、120ms 退场），支持毫秒级延迟防误触")
bullet("当前页面导航链接高亮（path startsWith 匹配）")
bullet("登录/注册页始终不显示导航栏（initNav 守卫逻辑）")

para("2. 深色模式", bold=True)
bullet("一键切换亮/暗主题，偏好持久化到 localStorage（ai-workhours-theme）")
bullet("50+ CSS 自定义属性（--bg-primary, --text-primary 等），覆盖背景、文字、边框、阴影所有色值")
bullet("0.35s cubic-bezier 过渡动画，切换平滑不闪屏")
bullet("图表实时跟随主题：绘制时通过 getComputedStyle 读取当前 CSS 变量颜色")

para("3. UI 组件（纯手写，无第三方 UI 库）", bold=True)
bullet("Toast 通知：右上角浮层，4 种类型（成功/失败/警告/信息），3.5 秒自动消失，点击提前关闭")
bullet("Modal 弹窗：毛玻璃遮罩层，入场缩放+淡入动画，点击遮罩关闭，用于删除确认/设置/编辑等场景")
bullet("Skeleton 骨架屏：加载态灰色脉冲动画占位，提升感知性能")
bullet("键盘快捷键：⌘1 仪表盘、⌘2 工时录入、⌘3 汇总统计、⌘K 切换主题（输入框聚焦时不触发）")
bullet("按钮 6 种样式：primary 填充蓝、outline 描边、danger 红、ghost 透明、icon 圆形、block 通栏")
bullet("标签 6 色：default/blue/green/red/orange/purple，对应不同业务场景")
bullet("表单控件：10px 圆角输入框、蓝色聚焦环（4px border + box-shadow）、14px 标签字体")

para("4. 响应式适配", bold=True)
bullet("768px 断点：导航高度缩减、表单行竖向堆叠、统计卡片 2 列、主网格单列")
bullet("480px 断点：导航链接间距缩小、字体缩减")
bullet("系统字体栈：SF Pro Display → PingFang SC → Helvetica Neue，各平台最优渲染")

para("5. API 调用封装", bold=True)
bullet("5 个通用函数：apiGet/apiPost/apiPut/apiDelete/apiPostDownload")
bullet("自动附加 Bearer Token、设置 Content-Type、解析 JSON 响应")
bullet("统一 401 处理：Token 过期/无效自动清空 localStorage 并跳转登录页")
bullet("统一错误处理：提取 FastAPI 的 detail 字段（兼容字符串和数组两种格式）")

doc.add_paragraph()

# ── 1.9 部署与运维 ──
heading("1.9 部署与运维", level=2)

bullet("Docker 容器化：基于 python:3.12-slim 镜像，Uvicorn 启动，8000 端口")
bullet("腾讯云 CloudBase CloudRun Serverless 部署：tcb cloudrun CLI 一键推送")
bullet("CORS 全局允许跨域（开发阶段），生产可按需收紧")
bullet("启动脚本（启动.command）：本地一键启动开发服务器，--reload 热重载")

doc.add_page_break()

# ═══════════════════════════════════════════
# PART 2: OPTIMIZATION OPPORTUNITIES
# ═══════════════════════════════════════════
heading("第二部分：可优化事项", level=1)

heading("2.1 安全相关（高优先级）", level=2)
make_table(
    ["序号", "问题", "影响", "建议"],
    [
        ["1", "SECRET_KEY 硬编码在 config.py", "源码泄露 → Token 可被伪造", "改为环境变量 SECRET_KEY"],
        ["2", "无登录速率限制", "可暴力破解密码", "同一 IP 5 分钟最多 10 次"],
        ["3", "Token 存 localStorage 有 XSS 风险", "恶意脚本可窃取 Token", "加 HttpOnly Cookie + CSRF Token"],
        ["4", "SQLite 用于生产环境", "并发写入受限，云端容器重建丢数据", "迁移至腾讯云 PostgreSQL"],
        ["5", "部分页面 HTML 拼接用户数据未转义", "潜在 XSS 注入", "统一加 escHtml() 过滤"],
    ],
    [1, 4, 3.5, 5.5]
)

heading("2.2 功能补全（中优先级）", level=2)
make_table(
    ["序号", "问题", "当前状态", "建议"],
    [
        ["6", "月总结页为空", "导航有入口但无功能", "实现月度工作总结增删改查 + 自动摘要"],
        ["7", "绩效评价页为空", "导航有入口但无功能", "考核模板 + 自评/管理员评分 + 总分计算"],
        ["8", "目标设定页为空", "导航有入口但无功能", "OKR 制定 + 进度跟踪 + 仪表盘嵌入"],
        ["9", "无密码找回", "忘记密码无法恢复", "邮箱验证 → 重置密码流程"],
        ["10", "无通知/提醒", "任务分配无感知", "企业微信机器人推送 + 站内消息"],
        ["11", "TaskUpdate Schema 定义了但没 API", "冗余代码未使用", "补任务编辑接口或清理 Schema"],
        ["12", "后端接收 assignee_ids 为空不校验", "可创建无执行人的任务", "加后端校验至少一位执行人"],
        ["13", "工时类型 8 个选项写死在 HTML 和 API", "无法自定义增减", "改为数据库驱动，管理员可配置"],
    ],
    [1, 3.5, 3.5, 5.5]
)

heading("2.3 性能与代码质量（中优先级）", level=2)
make_table(
    ["序号", "问题", "影响", "建议"],
    [
        ["14", "汇总在 Python 内存中用 for 循环聚合", "万级数据变慢", "改用 SQL GROUP BY + SUM"],
        ["15", "任务列表等查询存在 N+1 问题", "多次查库浪费性能", "统一加 joinedload 预加载关联"],
        ["16", "工时列表仅 limit 无 offset", "无法翻页，百条以上看不到", "加前端分页组件 + 后端 offset 参数"],
        ["17", "导出年份默认写死 2026", "2027 年需改代码", "取当前年份动态生成"],
        ["18", "零测试覆盖", "改代码靠人工验证", "加 pytest 核心流程测试（登录/录入/导出）"],
        ["19", "完成情况/项目编号为自由文本无约束", "同一项目多种写法", "加预设选项或联想补全"],
    ],
    [1, 3.5, 3.5, 5.5]
)

heading("2.4 体验优化（低优先级）", level=2)
make_table(
    ["序号", "问题", "建议"],
    [
        ["20", "无结构化日志 / 审计追踪", "logging 模块记录登录/删除/导出等关键操作"],
        ["21", "移动端未专门适配", "表单大按钮、表格横向滚动、底部 Tab 导航"],
        ["22", "快捷键提示在登录页显示但不生效", "登录页移除快捷键行，或在 app.js 中开放"],
        ["23", "项目重名无校验", "创建项目时检查同名项目是否存在"],
        ["24", "无数据导出格式选择", "加 CSV/PDF 导出选项"],
    ],
    [1, 5, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# PART 3: THIS WEEK PLAN
# ═══════════════════════════════════════════
heading("第三部分：本周计划（6月9日 — 6月13日）", level=1)

para("目标：补全三个存根页面 + 安全加固", bold=True, size=12)
doc.add_paragraph()

heading("3.1 月总结功能（周一 — 周二）", level=2)
bullet("数据库：新建 monthly_summaries 表（user_id, year_month, content, created_at, updated_at）")
bullet("后端：GET/POST/PUT 接口，按用户+月份唯一约束，自动从当月工时生成摘要草稿")
bullet("前端：月份选择器 + 富文本编辑区（textarea 即可），左侧已有总结列表")
bullet("仪表盘嵌入「本月总结状态」卡片（已写/未写）")

heading("3.2 绩效评价表（周三 — 周四上午）", level=2)
bullet("数据库：performance_templates（指标名、权重、评分标准）+ performance_scores（自评+管理员评）")
bullet("后端：管理员 CRUD 模板，成员提交自评，管理员评分，自动加权计算总分")
bullet("前端：评分卡片式界面，各项指标滑动条或星级打分，实时预览总分")
bullet("管理员视图：按人查看所有评分，对比自评 vs 管理员评分")

heading("3.3 安全加固（周四下午 — 周五）", level=2)
bullet("config.py 的 SECRET_KEY 改为 os.environ.get('SECRET_KEY', default_value)")
bullet("CloudRun 环境变量面板配置 SECRET_KEY")
bullet("登录接口加慢速限制：fastapi-limiter 或手动 IP 计数（同一 IP 5 分钟最多 10 次）")
bullet("SQLite → PostgreSQL 迁移预研：腾讯云数据库开服、连接串格式、SQLAlchemy 兼容性")
bullet("加 logging 模块：按日期切割日志文件，记录登录、删除工时、导出、任务操作")

doc.add_page_break()

# ═══════════════════════════════════════════
# PART 4: NEXT WEEK PLAN
# ═══════════════════════════════════════════
heading("第四部分：下周计划（6月16日 — 6月20日）", level=1)

para("目标：目标设定 + 数据库升级 + 体验打磨", bold=True, size=12)
doc.add_paragraph()

heading("4.1 目标设定表（周一 — 周二）", level=2)
bullet("数据库：goals 表（user_id, title, description, key_results JSON, progress 0-100, deadline, status）")
bullet("后端：目标 CRUD + 进度百分比更新 + 按状态/截止日期筛选")
bullet("前端：OKR 卡片布局，进度环/进度条，目标层级（个人目标 → 团队目标）")
bullet("仪表盘嵌入「本月目标进度」卡片，显示最紧急 3 个目标")
bullet("过期未完成目标红色高亮提醒")

heading("4.2 数据库升级 PostgreSQL（周三）", level=2)
bullet("腾讯云 PostgreSQL 实例开通（最低配即可，开发阶段够用）")
bullet("导出 SQLite 数据 → 转换为 PostgreSQL 兼容格式 → 导入")
bullet("修改 SQLAlchemy create_engine 连接串（支持环境变量 DATABASE_URL）")
bullet("CloudRun 配置 DATABASE_URL 环境变量 + 私有网络打通")
bullet("全量接口回归测试，确保 SQLAlchemy 方言兼容")

heading("4.3 通知与体验优化（周四 — 周五）", level=2)
bullet("企业微信机器人 Webhook 通知：任务分配时推送「XX 给你分配了任务：{title}」")
bullet("工时类型可配置：admin 后台增删工时类型，前端下拉动态获取")
bullet("分页：工时列表 + 任务列表加前端分页条 + 后端 offset/limit 参数")
bullet("移动端触摸优化：增大按钮热区（min 44px）、表格 overflow-x 滚动、底部固定导航")
bullet("代码清理：移除 TaskUpdate 未使用 Schema，所有查询统一 joinedload 策略")
bullet("加 5-10 条核心 pytest 用例：注册→登录→录入工时→查询→导出→删除")

doc.add_paragraph()

para("")
para("陈小黑，以上是完整的两周计划，你看完有调整随时说。", bold=False, size=12)

# ── save ──
output_path = "/Users/chen/Desktop/ai-work-hours-system_2/周报计划_6月9日.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
