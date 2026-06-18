"""Generate tech stack report as Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── global style defaults ──
style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    return p


def set_cell(cell, text, bold=False, center=True, font_size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.bold = bold
    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "333333")
        borders.append(el)
    tblPr.append(borders)


def shade_row(row, color="4472C4"):
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)


# ═══════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI工时管理系统\n技术栈报告")
run.font.size = Pt(26)
run.bold = True
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("2026年6月  ·  版本 1.0.0")
run.font.size = Pt(13)
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. 项目概述
# ═══════════════════════════════════════════
heading("一、项目概述", level=1)
para("AI工时管理系统是一个面向运维团队的 Web 应用，支持工时填报、月度汇总导出、任务派发与跟踪、绩效管理等日常管理功能。整体采用前后端一体的轻量架构，部署在腾讯云 CloudBase CloudRun（Serverless 容器服务）上。")

# ═══════════════════════════════════════════
# 2. 后端技术
# ═══════════════════════════════════════════
heading("二、后端技术", level=1)

heading("2.1 Python 生态", level=2)
para("后端开发语言为 Python 3.12，核心依赖如下：")

backend_table = doc.add_table(rows=10, cols=3)
backend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_border(backend_table)

headers = ["技术库", "版本", "用途"]
for i, h in enumerate(headers):
    set_cell(backend_table.rows[0].cells[i], h, bold=True, font_size=10)
shade_row(backend_table.rows[0], "4472C4")
for cell in backend_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data = [
    ["FastAPI", "0.104.1", "高性能 Web 框架，提供 REST API 和页面路由，支持自动生成 OpenAPI 文档"],
    ["Uvicorn", "0.48.0", "ASGI 服务器，用于运行 FastAPI 应用"],
    ["SQLAlchemy", "2.0.50", "ORM（对象关系映射），用 Python 对象操作数据库，避免手写 SQL"],
    ["Pydantic", "2.13.4", "请求/响应数据校验，自动类型检查与序列化"],
    ["python-jose", "3.5.0", "JWT（JSON Web Token）令牌生成与验证，用户认证的核心"],
    ["passlib + bcrypt", "1.7.4 / 3.2.2", "密码哈希加密，bcrypt 算法保证密码不可逆"],
    ["Jinja2", "3.1.2", "服务端模板引擎，将数据渲染为 HTML 页面"],
    ["openpyxl", "3.1.5", "Excel 文件读写，用于模板填充与工时导出"],
    ["python-multipart", "0.0.29", "表单/文件上传数据解析"],
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        set_cell(backend_table.rows[r].cells[c], val, font_size=9)

doc.add_paragraph()

heading("2.2 数据库", level=2)
para("数据库采用 SQLite（通过 SQLAlchemy 驱动），数据库文件 workhours.db 存储在项目根目录。")
para("数据表结构：", bold=True)

db_table = doc.add_table(rows=7, cols=3)
db_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_border(db_table)

for i, h in enumerate(["表名", "说明", "核心字段"]):
    set_cell(db_table.rows[0].cells[i], h, bold=True, font_size=10)
shade_row(db_table.rows[0], "4472C4")
for cell in db_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

db_data = [
    ["users", "用户表", "username, password_hash, real_name, role (admin/manager/member)"],
    ["projects", "项目表", "name, category (方舱/运维项目), description"],
    ["work_entries", "工时记录表", "user_id, project_id, hours, work_type, entry_date, location, project_number, completion_status"],
    ["user_sessions", "登录会话表", "user_id, token, device_type, is_active (支持多端登录)"],
    ["tasks", "任务表", "publisher_id, project_id, title, status (pending/in_progress/completed)"],
    ["task_assignments", "任务分配表", "task_id, user_id, status (pending/accepted/rejected/done)"],
]
for r, row_data in enumerate(db_data, 1):
    for c, val in enumerate(row_data):
        set_cell(db_table.rows[r].cells[c], val, font_size=9)

doc.add_paragraph()

heading("2.3 认证机制", level=2)
para("基于 JWT（JSON Web Token）的无状态认证方案：")
para("- 用户密码使用 bcrypt 算法哈希后存储，无法反向破解")
para("- 登录成功后生成 JWT Token，有效期 24 小时")
para("- 支持多端同时登录，每个设备独立会话")
para("- API 请求通过 Authorization: Bearer <token> 请求头验证身份")

# ═══════════════════════════════════════════
# 3. 前端技术
# ═══════════════════════════════════════════
heading("三、前端技术", level=1)

heading("3.1 技术选型", level=2)
para("前端采用原生 HTML + CSS + JavaScript，不依赖任何前端框架（React/Vue/Angular）。页面通过 Jinja2 模板引擎在服务端渲染后发送到浏览器，SPA 风格的交互（fetch API + DOM 操作）由原生 JS 实现。")

heading("3.2 页面路由", level=2)

page_table = doc.add_table(rows=11, cols=3)
page_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_border(page_table)

for i, h in enumerate(["路由", "页面", "功能说明"]):
    set_cell(page_table.rows[0].cells[i], h, bold=True, font_size=10)
shade_row(page_table.rows[0], "4472C4")
for cell in page_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

pages = [
    ["/", "login.html", "用户登录"],
    ["/register", "register.html", "账号注册"],
    ["/dashboard", "dashboard.html", "仪表盘，总工时、项目分布等概览统计"],
    ["/entry", "entry_form.html", "工时录入表单 + 近期记录列表"],
    ["/summary", "summary.html", "按日/周/月汇总统计 + Excel 模板导出"],
    ["/monthly-summary", "monthly_summary.html", "月度工作总结与复盘"],
    ["/tasks", "tasks.html", "任务发布、分配、状态跟踪"],
    ["/admin", "admin.html", "管理后台（用户管理、项目管理、会话管理）"],
    ["/users", "users.html", "人员信息查看与筛选"],
    ["/performance", "performance.html", "绩效评价表（考核指标与评分）"],
]
for r, row_data in enumerate(pages, 1):
    for c, val in enumerate(row_data):
        set_cell(page_table.rows[r].cells[c], val, font_size=9)

doc.add_paragraph()

heading("3.3 UI 组件", level=2)
para("全部 UI 组件为自实现，不依赖 Bootstrap 等第三方 CSS 框架：")
para("- Toast 浮层通知")
para("- Modal 模态弹窗（确认删除、表单提交等）")
para("- Skeleton 骨架屏加载占位")
para("- 下拉菜单（导航栏、用户设置）")
para("- 浅色主题（CSS 自定义属性 / data-theme）")

# ═══════════════════════════════════════════
# 4. 部署架构
# ═══════════════════════════════════════════
heading("四、部署架构", level=1)

heading("4.1 容器化", level=2)
para("通过 Dockerfile 构建容器镜像，基于 python:3.12-slim 基础镜像，安装依赖后由 Uvicorn 启动 FastAPI 应用，监听 8000 端口。")
para("启动命令：uvicorn main:app --host 0.0.0.0 --port 8000")

heading("4.2 云服务", level=2)
para("部署平台：腾讯云 CloudBase CloudRun")
para("- 类型：Serverless 容器服务（按需付费，自动扩缩容）")
para("- 特点：缩容到零后首个请求有 30~60 秒冷启动延迟")
para("- 环境 ID：tencnetcloud-d0gix13uz6b520764")
para("- 线上地址：https://workhours-263585-8-1317779850.sh.run.tcloudbase.com/")
para("- 部署方式：tcb cloudrun deploy CLI 命令")

heading("4.3 数据持久化", level=2)
para("重要区分：", bold=True)
para("- 本地开发：SQLite 文件（workhours.db）存储在项目目录 /Users/chen/Desktop/ai-work-hours-system_2/，数据持久保留，即使关机重启也不丢失")
para("- 云端部署：CloudRun 为无状态容器，文件系统是临时的（容器重建后数据丢失）。如需云端数据持久化，需接入腾讯云 MySQL 或 PostgreSQL 等云数据库")

# ═══════════════════════════════════════════
# 5. 架构特点
# ═══════════════════════════════════════════
heading("五、架构特点总结", level=1)

summary_table = doc.add_table(rows=8, cols=2)
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_border(summary_table)

for i, h in enumerate(["维度", "说明"]):
    set_cell(summary_table.rows[0].cells[i], h, bold=True, font_size=10)
shade_row(summary_table.rows[0], "4472C4")
for cell in summary_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

summary = [
    ["架构模式", "前后端一体（单体应用），服务端 Jinja2 渲染 + 前端 SPA 交互混合"],
    ["数据库", "本地 SQLite（开发环境），云端需升级为云数据库以保障数据持久化"],
    ["认证", "JWT 无状态认证 + bcrypt 密码哈希，支持多端会话"],
    ["Excel 导出", "openpyxl 模板填充，支持公式保留、自动行扩充、按用户+项目分组"],
    ["部署方式", "Docker 容器镜像 → CloudBase CloudRun Serverless"],
    ["前端技术", "无第三方前端框架，纯原生 JS + CSS + Jinja2 模板"],
    ["代码规模", "核心约 1000 行 Python（main.py 400 行 + crud.py 320 行 + export_utils.py 240 行）"],
]
for r, row_data in enumerate(summary, 1):
    for c, val in enumerate(row_data):
        set_cell(summary_table.rows[r].cells[c], val, bold=(c == 0), font_size=9)

doc.add_paragraph()

# ═══════════════════════════════════════════
# 6. 架构图
# ═══════════════════════════════════════════
heading("六、技术架构图", level=1)

arch_text = """┌──────────────────────────────────────────────────────┐
│                      浏览器                            │
│           HTML + CSS + JavaScript                      │
│       (Jinja2 模板渲染, fetch API 交互)                  │
└────────────────────┬─────────────────────────────────┘
                     │  HTTP / REST
┌────────────────────▼─────────────────────────────────┐
│                FastAPI (Python 3.12)                    │
│                                                        │
│  ┌──────────┐  ┌──────────┐  ┌──────────────────┐     │
│  │ 路由系统  │  │ 认证中间件│  │  模板引擎 (Jinja2) │     │
│  │ (REST API│  │  (JWT)   │  │                   │     │
│  │ + Pages) │  │          │  │                   │     │
│  └──────────┘  └──────────┘  └──────────────────┘     │
│          │                    │                         │
│  ┌───────▼────────────────────▼───────────────┐       │
│  │          SQLAlchemy ORM                     │       │
│  │      (Pydantic 数据校验 + CRUD)             │       │
│  └────────────────────┬───────────────────────┘       │
│                       │                                │
│  ┌────────────────────▼───────────────────────┐       │
│  │           SQLite 数据库                     │       │
│  │          (workhours.db)                     │       │
│  └────────────────────────────────────────────┘       │
│                                                        │
│  ┌────────────────────────────────────────────┐       │
│  │    openpyxl Excel 模板导出                    │       │
│  │   (月度工时报 / 公式保留 / 自动行扩展)          │       │
│  └────────────────────────────────────────────┘       │
│                                                        │
│               Uvicorn ASGI Server                       │
└────────────────────┬─────────────────────────────────┘
                     │
┌────────────────────▼─────────────────────────────────┐
│      Docker 容器 (python:3.12-slim)                     │
│      CloudBase CloudRun Serverless                      │
└──────────────────────────────────────────────────────┘"""

p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = "Courier New"
run.font.size = Pt(7.5)

# ── save ──
output_path = "/Users/chen/Desktop/ai-work-hours-system_2/技术栈报告.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
